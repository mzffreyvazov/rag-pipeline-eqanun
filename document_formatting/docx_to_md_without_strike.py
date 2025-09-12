#!/usr/bin/env python3
"""
DOCX → Markdown (skip strikethrough)

Reads .docx files and writes Markdown while omitting any text runs
with strikethrough formatting. Other formatting (bold, italic,
headings, lists, tables, alignment) is preserved similarly to the
existing docx_to_md.py.

Requirements:
    pip install python-docx

Usage:
    - Edit the `input_files` list and optional `output_dir` in main().
    - Run: python docx_to_md_without_strike.py
"""

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.text.run import Run
import re
import sys
import os


def has_strikethrough(run) -> bool:
    # python-docx sets run.font.strike to True/False/None
    return getattr(run.font, "strike", None) is True


class DocxToMarkdownNoStrike:
    def __init__(self):
        self._parts = []

    # --- Runs / inline formatting ---
    def run_to_md(self, run: Run) -> str:
        if has_strikethrough(run):
            return ""  # skip strikethrough text entirely

        text = run.text or ""
        if not text:
            return text

        # bold / italic (preserve like the original converter)
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"

        # superscript / subscript via HTML (Markdown has no native syntax)
        # Note: If both are somehow set, prefer superscript.
        if getattr(run.font, "superscript", None):
            text = f"<sup>{text}</sup>"
        elif getattr(run.font, "subscript", None):
            text = f"<sub>{text}</sub>"

        return text

    # --- Paragraph-level helpers ---
    def heading_level(self, paragraph) -> int:
        style_name = (paragraph.style.name or "").lower()
        if "heading" in style_name:
            nums = re.findall(r"\d+", style_name)
            if nums:
                return min(int(nums[0]), 6)
        return 0

    def paragraph_to_md(self, paragraph) -> str:
        lvl = self.heading_level(paragraph)

        # Render inline runs (skipping strikethrough ones)
        buf = []
        for run in paragraph.runs:
            part = self.run_to_md(run)
            if part:
                buf.append(part)
        content = "".join(buf).strip()

        # If paragraph ends up empty (e.g., everything was strikethrough), skip it
        if not content:
            return ""

        if lvl > 0:
            return f"{'#' * lvl} {content}\n"

        # Simple list detection (style or leading bullet/number)
        style_name = paragraph.style.name or ""
        if style_name.startswith("List"):
            if "Bullet" in style_name:
                return f"- {content}\n"
            else:
                return f"1. {content}\n"
        else:
            # Also handle literal bullets / numbers at start of original text
            raw = paragraph.text.strip()
            if re.match(r"^[•\-\*]\s+", raw):
                return f"- {content}\n"
            if re.match(r"^\d+\.\s+", raw):
                return f"1. {content}\n"

        # Alignment handling for normal paragraphs
        align = paragraph.alignment
        if align == WD_PARAGRAPH_ALIGNMENT.CENTER:
            return f"<center>{content}</center>\n\n"
        if align == WD_PARAGRAPH_ALIGNMENT.RIGHT:
            return f"<div align=\"right\">{content}</div>\n\n"

        return f"{content}\n\n"

    # --- Tables ---
    def table_to_md(self, table) -> str:
        rows_md = []
        for r_idx, row in enumerate(table.rows):
            cols = []
            for cell in row.cells:
                cell_buf = []
                for p in cell.paragraphs:
                    # render paragraph content (skip if empty)
                    p_md = self.paragraph_to_md(p).strip()
                    if p_md:
                        # strip trailing md line breaks for table cell
                        cell_buf.append(p_md.replace("\n", " "))
                cols.append(" ".join(cell_buf).strip())

            rows_md.append("| " + " | ".join(cols) + " |")
            if r_idx == 0:
                rows_md.append("| " + " | ".join(["---"] * len(cols)) + " |")
        return "\n".join(rows_md) + "\n\n"

    # --- Document traversal preserving body order ---
    def convert(self, doc: Document) -> str:
        parts = []
        body = doc.element.body
        for el in body:
            tag = el.tag.rsplit('}', 1)[-1]
            if tag == 'p':
                # find matching paragraph object by element identity
                for para in doc.paragraphs:
                    if para._element is el:
                        md = self.paragraph_to_md(para)
                        if md:
                            parts.append(md)
                        break
            elif tag == 'tbl':
                for tbl in doc.tables:
                    if tbl._element is el:
                        parts.append(self.table_to_md(tbl))
                        break
        # condense blank lines
        out = "".join(parts)
        out = re.sub(r"\n{3,}", "\n\n", out).strip() + "\n"
        return out


# --- Top-level API ---
def convert_docx_to_markdown_without_strike(input_path: str, output_path: str) -> bool:
    try:
        print(f"Loading document: {input_path}")
        doc = Document(input_path)

        print("Converting to Markdown (skipping strikethrough)...")
        conv = DocxToMarkdownNoStrike()
        md = conv.convert(doc)

        print(f"Saving Markdown: {output_path}")
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(md)
        print("Done.")
        return True
    except FileNotFoundError:
        print(f"Error: Input file '{input_path}' not found.")
        return False
    except PermissionError:
        print("Error: Permission denied. Check file permissions.")
        return False
    except Exception as e:
        print(f"Error converting document: {e}")
        return False


def main():
    # Configure inputs here: a list of .docx file paths to convert
    input_files = [
        # Examples:
        # r"D:\\Downloads\\code\\code\\Extra-Projects\\rag-pipeline-eqanun\\document_formatting\\mecelleler-docx\\output-cinayet.docx",
        # r"D:\\path\\to\\another.docx",
    ]

    # Optional output directory. If None, writes next to each input file.
    output_dir = "D:\\Downloads\\code\\code\\Extra-Projects\\rag-pipeline-eqanun\\document_formatting\\mecelleler-raw\\mecelleler-wo-strikes" # e.g., r"D:\\Downloads\\...\\document_formatting\\outputs_md"

    if not input_files:
        print("No input files configured. Edit 'input_files' in this script.")
        return

    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    for input_file in input_files:
        if not os.path.exists(input_file):
            print(f"Skip: '{input_file}' not found.")
            continue

        if not input_file.lower().endswith('.docx'):
            print(f"Warning: '{input_file}' doesn't have .docx extension. Proceeding anyway...")

        base = os.path.splitext(os.path.basename(input_file))[0]
        out_dir = output_dir if output_dir else os.path.dirname(input_file)
        output_file = os.path.join(out_dir, base + '.md')

        ok = convert_docx_to_markdown_without_strike(input_file, output_file)
        if ok:
            print(f"Converted: '{input_file}' -> '{output_file}' (strikethrough omitted)")
        else:
            print(f"Failed: '{input_file}'")


if __name__ == "__main__":
    main()
