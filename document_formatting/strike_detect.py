#!/usr/bin/env python3
"""
Word Document Strikethrough Text Detector

This script processes Word documents (.docx) and identifies text with strikethrough formatting.
It can process a single file or multiple files in a directory.

Requirements:
    pip install python-docx

Usage:
    python strikethrough_detector.py path/to/document.docx
    python strikethrough_detector.py path/to/directory/
"""

import os
from pathlib import Path
from typing import List, Tuple, Dict
from docx import Document
from docx.shared import RGBColor

import json


class StrikethroughDetector:
    def __init__(self):
        self.results = []
    
    def detect_strikethrough_in_document(self, file_path: str) -> List[Dict]:
        """
        Detect strikethrough text in a Word document.
        
        Args:
            file_path (str): Path to the Word document
            
        Returns:
            List[Dict]: List of dictionaries containing strikethrough text information
        """
        try:
            doc = Document(file_path)
            strikethrough_items = []
            
            # Process paragraphs
            for para_idx, paragraph in enumerate(doc.paragraphs):
                para_strikethrough = self._find_strikethrough_in_paragraph(paragraph, para_idx)
                strikethrough_items.extend(para_strikethrough)
            
            # Process tables
            for table_idx, table in enumerate(doc.tables):
                table_strikethrough = self._find_strikethrough_in_table(table, table_idx)
                strikethrough_items.extend(table_strikethrough)
            
            return strikethrough_items
            
        except Exception as e:
            print(f"Error processing {file_path}: {str(e)}")
            return []
    
    def _find_strikethrough_in_paragraph(self, paragraph, para_idx: int) -> List[Dict]:
        """Find strikethrough text in a paragraph."""
        strikethrough_items = []
        
        for run_idx, run in enumerate(paragraph.runs):
            if run.font.strike or run.font.double_strike:
                strikethrough_items.append({
                    'type': 'paragraph',
                    'location': f'Paragraph {para_idx + 1}, Run {run_idx + 1}',
                    'text': run.text.strip(),
                    'strike_type': 'double' if run.font.double_strike else 'single',
                    'paragraph_text': paragraph.text.strip()
                })
        
        return strikethrough_items
    
    def _find_strikethrough_in_table(self, table, table_idx: int) -> List[Dict]:
        """Find strikethrough text in table cells."""
        strikethrough_items = []
        
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    for run_idx, run in enumerate(paragraph.runs):
                        if run.font.strike or run.font.double_strike:
                            strikethrough_items.append({
                                'type': 'table',
                                'location': f'Table {table_idx + 1}, Row {row_idx + 1}, Cell {cell_idx + 1}, Paragraph {para_idx + 1}, Run {run_idx + 1}',
                                'text': run.text.strip(),
                                'strike_type': 'double' if run.font.double_strike else 'single',
                                'cell_text': cell.text.strip()
                            })
        
        return strikethrough_items
    
    def process_file(self, file_path: str) -> Dict:
        """Process a single Word document file."""
        print(f"Processing: {file_path}")
        
        strikethrough_items = self.detect_strikethrough_in_document(file_path)
        
        result = {
            'file_path': file_path,
            'filename': os.path.basename(file_path),
            'strikethrough_count': len(strikethrough_items),
            'strikethrough_items': strikethrough_items
        }
        
        return result
    
    def process_directory(self, directory_path: str) -> List[Dict]:
        """Process all Word documents in a directory."""
        results = []
        directory = Path(directory_path)
        
        # Find all .docx files
        docx_files = list(directory.glob('*.docx'))
        
        if not docx_files:
            print(f"No .docx files found in {directory_path}")
            return results
        
        print(f"Found {len(docx_files)} Word document(s) to process")
        
        for file_path in docx_files:
            # Skip temporary files (starting with ~$)
            if file_path.name.startswith('~$'):
                continue
                
            result = self.process_file(str(file_path))
            results.append(result)
        
        return results
    
    def print_results(self, results: List[Dict]):
        """Print the results in a readable format."""
        total_files = len(results)
        total_strikethrough = sum(result['strikethrough_count'] for result in results)
        
        print("\n" + "="*60)
        print("STRIKETHROUGH DETECTION RESULTS")
        print("="*60)
        print(f"Files processed: {total_files}")
        print(f"Total strikethrough instances found: {total_strikethrough}")
        print("="*60)
        
        for result in results:
            print(f"\nFile: {result['filename']}")
            print(f"Path: {result['file_path']}")
            print(f"Strikethrough instances: {result['strikethrough_count']}")
            
            if result['strikethrough_items']:
                print("\nStrikethrough text found:")
                print("-" * 40)
                
                for idx, item in enumerate(result['strikethrough_items'], 1):
                    print(f"{idx}. Location: {item['location']}")
                    print(f"   Text: '{item['text']}'")
                    print(f"   Type: {item['strike_type']} strikethrough")
                    
                    if item['type'] == 'paragraph':
                        print(f"   Full paragraph: '{item['paragraph_text'][:100]}{'...' if len(item['paragraph_text']) > 100 else ''}'")
                    else:
                        print(f"   Full cell: '{item['cell_text'][:100]}{'...' if len(item['cell_text']) > 100 else ''}'")
                    print()
            else:
                print("   No strikethrough text found.")
            
            print("-" * 60)
    
    def save_results_to_file(self, results: List[Dict], output_file: str = "strikethrough_report.txt"):
        """Save results to a text file."""
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write("STRIKETHROUGH DETECTION REPORT\n")
                f.write("=" * 50 + "\n\n")
                
                total_files = len(results)
                total_strikethrough = sum(result['strikethrough_count'] for result in results)
                
                f.write(f"Files processed: {total_files}\n")
                f.write(f"Total strikethrough instances: {total_strikethrough}\n")
                f.write("=" * 50 + "\n\n")
                
                for result in results:
                    f.write(f"File: {result['filename']}\n")
                    f.write(f"Path: {result['file_path']}\n")
                    f.write(f"Strikethrough instances: {result['strikethrough_count']}\n\n")
                    
                    if result['strikethrough_items']:
                        f.write("Strikethrough text found:\n")
                        f.write("-" * 30 + "\n")
                        
                        for idx, item in enumerate(result['strikethrough_items'], 1):
                            f.write(f"{idx}. Location: {item['location']}\n")
                            f.write(f"   Text: '{item['text']}'\n")
                            f.write(f"   Type: {item['strike_type']} strikethrough\n")
                            
                            if item['type'] == 'paragraph':
                                f.write(f"   Full paragraph: '{item['paragraph_text']}'\n")
                            else:
                                f.write(f"   Full cell: '{item['cell_text']}'\n")
                            f.write("\n")
                    else:
                        f.write("No strikethrough text found.\n")
                    
                    f.write("-" * 50 + "\n\n")
            
            print(f"\nResults saved to: {output_file}")
            
        except Exception as e:
            print(f"Error saving results to file: {str(e)}")
    
    def get_results_as_json(self, results: List[Dict]) -> str:
        """Convert results to JSON format."""
        return json.dumps(results, ensure_ascii=False, indent=4)


def main():
    """Main function to run the strikethrough detector."""
    # Specify the file path directly in the code
    FILE_PATH = r"D:/Downloads/code/code/Extra-Projects/rag-pipeline-eqanun/document_formatting/mecelleler-docx/-Azərbaycan Respublikasının Mülki Prosessual Məcəlləsi.docx"
    
    # Extract filename for output file naming
    filename = os.path.basename(FILE_PATH).replace('.docx', '').replace(' ', '_').replace('-', '_')
    
    detector = StrikethroughDetector()
    
    try:
        # Process the document
        result = detector.process_file(FILE_PATH)
        
        # Convert to JSON and print
        json_output = detector.get_results_as_json([result])
        print("\nJSON Report:")
        print("=" * 60)
        print(json_output)
        
        # Save JSON to strikes_json folder in parent directory
        parent_dir = Path(__file__).parent
        strikes_json_dir = parent_dir / "strikes_json"
        strikes_json_dir.mkdir(exist_ok=True)
        
        output_file = strikes_json_dir / f"{filename}_strikethrough_report.json"
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(json_output)
        print(f"\nResults saved to: {output_file}")
        
        return json_output
        
    except Exception as e:
        error_result = {
            'error': str(e),
            'file_path': FILE_PATH,
            'filename': os.path.basename(FILE_PATH),
            'strikethrough_count': 0,
            'strikethrough_items': []
        }
        json_output = detector.get_results_as_json([error_result])
        print(f"Error occurred: {json_output}")
        
        # Save error JSON to strikes_json folder in parent directory
        parent_dir = Path(__file__).parent
        strikes_json_dir = parent_dir / "strikes_json"
        strikes_json_dir.mkdir(exist_ok=True)
        
        output_file = strikes_json_dir / f"{filename}_strikethrough_report.json"
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(json_output)
        print(f"\nError results saved to: {output_file}")
        
        return json_output


if __name__ == "__main__":
    main()