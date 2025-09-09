#!/usr/bin/env python3
"""
Word Document Strikethrough Text Detector

This script processes Word documents (.docx) and identifies text with strikethrough formatting.
It can process a single file or multiple files in a directory.

Requirements:
    pip install python-docx

Usage:
    python strikethrough_detector.py path/to/document.docx
    python strikethrough_detector.py path/to/directory/
"""

import os
import sys
from pathlib import Path
from typing import List, Tuple, Dict
from docx import Document
from docx.shared import RGBColor


class StrikethroughDetector:
    def __init__(self):
        self.results = []
    
    def detect_strikethrough_in_document(self, file_path: str) -> List[Dict]:
        """
        Detect strikethrough text in a Word document.
        
        Args:
            file_path (str): Path to the Word document
            
        Returns:
            List[Dict]: List of dictionaries containing strikethrough text information
        """
        try:
            doc = Document(file_path)
            strikethrough_items = []
            
            # Process paragraphs
            for para_idx, paragraph in enumerate(doc.paragraphs):
                para_strikethrough = self._find_strikethrough_in_paragraph(paragraph, para_idx)
                strikethrough_items.extend(para_strikethrough)
            
            # Process tables
            for table_idx, table in enumerate(doc.tables):
                table_strikethrough = self._find_strikethrough_in_table(table, table_idx)
                strikethrough_items.extend(table_strikethrough)
            
            return strikethrough_items
            
        except Exception as e:
            print(f"Error processing {file_path}: {str(e)}")
            return []
    
    def _find_strikethrough_in_paragraph(self, paragraph, para_idx: int) -> List[Dict]:
        """Find strikethrough text in a paragraph."""
        strikethrough_items = []
        
        for run_idx, run in enumerate(paragraph.runs):
            if run.font.strike or run.font.double_strike:
                strikethrough_items.append({
                    'type': 'paragraph',
                    'location': f'Paragraph {para_idx + 1}, Run {run_idx + 1}',
                    'text': run.text.strip(),
                    'strike_type': 'double' if run.font.double_strike else 'single',
                    'paragraph_text': paragraph.text.strip()
                })
        
        return strikethrough_items
    
    def _find_strikethrough_in_table(self, table, table_idx: int) -> List[Dict]:
        """Find strikethrough text in table cells."""
        strikethrough_items = []
        
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    for run_idx, run in enumerate(paragraph.runs):
                        if run.font.strike or run.font.double_strike:
                            strikethrough_items.append({
                                'type': 'table',
                                'location': f'Table {table_idx + 1}, Row {row_idx + 1}, Cell {cell_idx + 1}, Paragraph {para_idx + 1}, Run {run_idx + 1}',
                                'text': run.text.strip(),
                                'strike_type': 'double' if run.font.double_strike else 'single',
                                'cell_text': cell.text.strip()
                            })
        
        return strikethrough_items
    
    def process_file(self, file_path: str) -> Dict:
        """Process a single Word document file."""
        print(f"Processing: {file_path}")
        
        strikethrough_items = self.detect_strikethrough_in_document(file_path)
        
        result = {
            'file_path': file_path,
            'filename': os.path.basename(file_path),
            'strikethrough_count': len(strikethrough_items),
            'strikethrough_items': strikethrough_items
        }
        
        return result
    
    def process_directory(self, directory_path: str) -> List[Dict]:
        """Process all Word documents in a directory."""
        results = []
        directory = Path(directory_path)
        
        # Find all .docx files
        docx_files = list(directory.glob('*.docx'))
        
        if not docx_files:
            print(f"No .docx files found in {directory_path}")
            return results
        
        print(f"Found {len(docx_files)} Word document(s) to process")
        
        for file_path in docx_files:
            # Skip temporary files (starting with ~$)
            if file_path.name.startswith('~$'):
                continue
                
            result = self.process_file(str(file_path))
            results.append(result)
        
        return results
    
    def print_results(self, results: List[Dict]):
        """Print the results in a readable format."""
        total_files = len(results)
        total_strikethrough = sum(result['strikethrough_count'] for result in results)
        
        print("\n" + "="*60)
        print("STRIKETHROUGH DETECTION RESULTS")
        print("="*60)
        print(f"Files processed: {total_files}")
        print(f"Total strikethrough instances found: {total_strikethrough}")
        print("="*60)
        
        for result in results:
            print(f"\nFile: {result['filename']}")
            print(f"Path: {result['file_path']}")
            print(f"Strikethrough instances: {result['strikethrough_count']}")
            
            if result['strikethrough_items']:
                print("\nStrikethrough text found:")
                print("-" * 40)
                
                for idx, item in enumerate(result['strikethrough_items'], 1):
                    print(f"{idx}. Location: {item['location']}")
                    print(f"   Text: '{item['text']}'")
                    print(f"   Type: {item['strike_type']} strikethrough")
                    
                    if item['type'] == 'paragraph':
                        print(f"   Full paragraph: '{item['paragraph_text'][:100]}{'...' if len(item['paragraph_text']) > 100 else ''}'")
                    else:
                        print(f"   Full cell: '{item['cell_text'][:100]}{'...' if len(item['cell_text']) > 100 else ''}'")
                    print()
            else:
                print("   No strikethrough text found.")
            
            print("-" * 60)
    
    def save_results_to_file(self, results: List[Dict], output_file: str = "strikethrough_report.txt"):
        """Save results to a text file."""
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write("STRIKETHROUGH DETECTION REPORT\n")
                f.write("=" * 50 + "\n\n")
                
                total_files = len(results)
                total_strikethrough = sum(result['strikethrough_count'] for result in results)
                
                f.write(f"Files processed: {total_files}\n")
                f.write(f"Total strikethrough instances: {total_strikethrough}\n")
                f.write("=" * 50 + "\n\n")
                
                for result in results:
                    f.write(f"File: {result['filename']}\n")
                    f.write(f"Path: {result['file_path']}\n")
                    f.write(f"Strikethrough instances: {result['strikethrough_count']}\n\n")
                    
                    if result['strikethrough_items']:
                        f.write("Strikethrough text found:\n")
                        f.write("-" * 30 + "\n")
                        
                        for idx, item in enumerate(result['strikethrough_items'], 1):
                            f.write(f"{idx}. Location: {item['location']}\n")
                            f.write(f"   Text: '{item['text']}'\n")
                            f.write(f"   Type: {item['strike_type']} strikethrough\n")
                            
                            if item['type'] == 'paragraph':
                                f.write(f"   Full paragraph: '{item['paragraph_text']}'\n")
                            else:
                                f.write(f"   Full cell: '{item['cell_text']}'\n")
                            f.write("\n")
                    else:
                        f.write("No strikethrough text found.\n")
                    
                    f.write("-" * 50 + "\n\n")
            
            print(f"\nResults saved to: {output_file}")
            
        except Exception as e:
            print(f"Error saving results to file: {str(e)}")


def main():
    """Main function to run the strikethrough detector."""
    if len(sys.argv) != 2:
        print("Usage: python strikethrough_detector.py <file_or_directory_path>")
        print("Example: python strikethrough_detector.py document.docx")
        print("Example: python strikethrough_detector.py /path/to/documents/")
        sys.exit(1)
    
    input_path = sys.argv[1]
    
    if not os.path.exists(input_path):
        print(f"Error: Path '{input_path}' does not exist.")
        sys.exit(1)
    
    detector = StrikethroughDetector()
    
    try:
        if os.path.isfile(input_path):
            # Process single file
            if not input_path.lower().endswith('.docx'):
                print("Error: Please provide a .docx file.")
                sys.exit(1)
            
            result = detector.process_file(input_path)
            results = [result]
        
        elif os.path.isdir(input_path):
            # Process directory
            results = detector.process_directory(input_path)
        
        else:
            print(f"Error: '{input_path}' is neither a file nor a directory.")
            sys.exit(1)
        
        # Display and save results
        detector.print_results(results)
        
        # Ask user if they want to save results to file
        save_to_file = input("\nSave results to file? (y/n): ").strip().lower()
        if save_to_file in ['y', 'yes']:
            output_filename = input("Enter output filename (press Enter for 'strikethrough_report.txt'): ").strip()
            if not output_filename:
                output_filename = "strikethrough_report.txt"
            detector.save_results_to_file(results, output_filename)
    
    except KeyboardInterrupt:
        print("\nOperation cancelled by user.")
    except Exception as e:
        print(f"An error occurred: {str(e)}")


def quick_test():
    """Quick test function for a specific file."""
    # Uncomment and modify the line below with your file path for quick testing
    file_path = "D:\\Downloads\\code\\code\\Extra-Projects\\rag-pipeline-eqanun\\assets\\-Azərbaycan Respublikasının Torpaq Məcəlləsi.docx"

    if not os.path.exists(file_path):
        print(f"Error: File '{file_path}' does not exist.")
        return
    
    detector = StrikethroughDetector()
    result = detector.process_file(file_path)
    detector.print_results([result])
    
    # Optionally save results
    detector.save_results_to_file([result], "torpaq_mecellesi_strikethrough_report.txt")


if __name__ == "__main__":
    # Uncomment the line below to run quick test instead of main()
    quick_test()
    # main()