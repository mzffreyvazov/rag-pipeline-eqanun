#!/usr/bin/env python3
"""
DOCX Strikethrough Remover

This script reads a DOCX file, removes all text with strikethrough formatting,
and saves the result as a new DOCX file while preserving all other formatting.

Requirements:
    pip install python-docx

Usage:
    Modify the input_file and output_file variables in the main() function,
    then run: python docx_strikethrough_remover.py
"""

from docx import Document
from docx.shared import Inches
import sys
import os

def has_strikethrough(run):
    """
    Check if a run has strikethrough formatting.
    
    Args:
        run: A run object from python-docx
        
    Returns:
        bool: True if the run has strikethrough formatting
    """
    return run.font.strike is True

def process_paragraph(paragraph):
    """
    Process a paragraph by removing runs with strikethrough formatting.
    
    Args:
        paragraph: A paragraph object from python-docx
    """
    # Work backwards through runs to avoid index issues when removing
    for i in range(len(paragraph.runs) - 1, -1, -1):
        run = paragraph.runs[i]
        if has_strikethrough(run):
            # Remove the run element from the paragraph
            paragraph._element.remove(run._element)

def process_table(table):
    """
    Process a table by removing strikethrough text from all cells.
    
    Args:
        table: A table object from python-docx
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                process_paragraph(paragraph)

def remove_strikethrough_from_docx(input_path, output_path):
    """
    Remove strikethrough text from a DOCX file and save the result.
    
    Args:
        input_path (str): Path to the input DOCX file
        output_path (str): Path for the output DOCX file
    """
    try:
        # Load the document
        print(f"Loading document: {input_path}")
        doc = Document(input_path)
        
        # Process all paragraphs in the document body
        print("Processing paragraphs...")
        for paragraph in doc.paragraphs:
            process_paragraph(paragraph)
        
        # Process all tables in the document
        print("Processing tables...")
        for table in doc.tables:
            process_table(table)
        
        # Process headers and footers
        print("Processing headers and footers...")
        for section in doc.sections:
            # Process headers
            if section.header:
                for paragraph in section.header.paragraphs:
                    process_paragraph(paragraph)
                for table in section.header.tables:
                    process_table(table)
            
            # Process footers
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    process_paragraph(paragraph)
                for table in section.footer.tables:
                    process_table(table)
        
        # Save the processed document
        print(f"Saving processed document: {output_path}")
        doc.save(output_path)
        print("Document processing completed successfully!")
        
    except FileNotFoundError:
        print(f"Error: Input file '{input_path}' not found.")
        return False
    except PermissionError:
        print(f"Error: Permission denied. Check if '{output_path}' is open in another application.")
        return False
    except Exception as e:
        print(f"Error processing document: {str(e)}")
        return False
    
    return True

def main():
    """
    Main function to handle command line arguments or use default file paths.
    """
    if len(sys.argv) == 3:
        input_file = sys.argv[1]
        output_file = sys.argv[2]
    elif len(sys.argv) == 1:
        # Default file paths if no arguments provided
        input_file = "input.docx"
        output_file = "output.docx"
        
        # Check if default input file exists
        if not os.path.exists(input_file):
            print("Usage: python docx_strikethrough_remover.py <input.docx> <output.docx>")
            print("Or place your DOCX file as 'input.docx' in the same directory")
            return
    else:
        print("Usage: python docx_strikethrough_remover.py <input.docx> <output.docx>")
        return
    
    # Validate input file exists
    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' does not exist.")
        return
    
    # Validate input file extension
    if not input_file.lower().endswith('.docx'):
        print("Warning: Input file doesn't have .docx extension. Proceeding anyway...")
    
    # Process the document
    success = remove_strikethrough_from_docx(input_file, output_file)
    
    if success:
        print(f"\nSuccessfully processed '{input_file}' -> '{output_file}'")
        print("All strikethrough text has been removed while preserving other formatting.")
    else:
        print("Failed to process the document.")

if __name__ == "__main__":
    main()