#!/usr/bin/env python3
"""
DOCX to Markdown Converter

This script converts a DOCX file to Markdown format while preserving
formatting like bold, italic, headers, lists, tables, and links.

Requirements:
    pip install python-docx

Usage:
    Modify the input_file and output_file variables in the main() function,
    then run: python docx_to_markdown_converter.py
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import os

class DocxToMarkdownConverter:
    """
    A class to convert DOCX documents to Markdown format.
    """
    
    def __init__(self):
        self.markdown_content = []
        self.in_list = False
        self.list_level = 0
    
    def convert_run_to_markdown(self, run):
        """
        Convert a single run to markdown with appropriate formatting.
        
        Args:
            run: A run object from python-docx
            
        Returns:
            str: Markdown formatted text
        """
        text = run.text
        
        # Skip empty runs
        if not text.strip():
            return text
        
        # Apply formatting
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"
        
        if run.font.strike:
            text = f"~~{text}~~"
        
        # Handle superscript and subscript (markdown doesn't support these natively, 
        # but we can use HTML tags)
        if run.font.superscript:
            text = f"<sup>{text}</sup>"
        elif run.font.subscript:
            text = f"<sub>{text}</sub>"
        
        return text
    
    def get_heading_level(self, paragraph):
        """
        Determine if a paragraph is a heading and return its level.
        
        Args:
            paragraph: A paragraph object from python-docx
            
        Returns:
            int: Heading level (1-6) or 0 if not a heading
        """
        style_name = paragraph.style.name.lower()
        
        if 'heading' in style_name:
            # Extract number from heading style
            numbers = re.findall(r'\d+', style_name)
            if numbers:
                level = int(numbers[0])
                return min(level, 6)  # Markdown supports up to 6 levels
        
        return 0
    
    def convert_paragraph_to_markdown(self, paragraph):
        """
        Convert a paragraph to markdown format.
        
        Args:
            paragraph: A paragraph object from python-docx
            
        Returns:
            str: Markdown formatted paragraph
        """
        # Check if it's a heading
        heading_level = self.get_heading_level(paragraph)
        
        # Convert all runs in the paragraph
        markdown_text = ""
        for run in paragraph.runs:
            markdown_text += self.convert_run_to_markdown(run)
        
        # Skip empty paragraphs
        if not markdown_text.strip():
            return ""
        
        # Handle headings
        if heading_level > 0:
            return f"{'#' * heading_level} {markdown_text.strip()}\n"
        
        # Handle list items
        if paragraph.style.name.startswith('List'):
            if 'Bullet' in paragraph.style.name or paragraph.text.strip().startswith(('•', '-', '*')):
                # Remove bullet characters if they exist
                text = re.sub(r'^[•\-\*]\s*', '', markdown_text.strip())
                return f"- {text}\n"
            else:
                # Numbered list - let markdown handle the numbering
                text = re.sub(r'^\d+\.\s*', '', markdown_text.strip())
                return f"1. {text}\n"
        
        # Handle alignment for regular paragraphs
        alignment = paragraph.alignment
        if alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            return f"<center>{markdown_text.strip()}</center>\n\n"
        elif alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
            return f"<div align=\"right\">{markdown_text.strip()}</div>\n\n"
        
        # Regular paragraph
        return f"{markdown_text.strip()}\n\n"
    
    def convert_table_to_markdown(self, table):
        """
        Convert a table to markdown format.
        
        Args:
            table: A table object from python-docx
            
        Returns:
            str: Markdown formatted table
        """
        markdown_table = []
        
        for i, row in enumerate(table.rows):
            row_text = []
            for cell in row.cells:
                # Get text from all paragraphs in the cell
                cell_text = ""
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        cell_text += self.convert_run_to_markdown(run)
                
                # Clean up cell text and handle newlines
                cell_text = cell_text.replace('\n', ' ').strip()
                row_text.append(cell_text)
            
            # Create table row
            markdown_table.append("| " + " | ".join(row_text) + " |")
            
            # Add header separator after first row
            if i == 0:
                separator = "| " + " | ".join(["---"] * len(row_text)) + " |"
                markdown_table.append(separator)
        
        return "\n".join(markdown_table) + "\n\n"
    
    def convert_document(self, doc):
        """
        Convert the entire document to markdown.
        
        Args:
            doc: A Document object from python-docx
            
        Returns:
            str: Complete markdown content
        """
        markdown_parts = []
        
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                # Find the corresponding paragraph object
                for para in doc.paragraphs:
                    if para._element == element:
                        markdown_para = self.convert_paragraph_to_markdown(para)
                        if markdown_para:
                            markdown_parts.append(markdown_para)
                        break
            
            elif element.tag.endswith('tbl'):  # Table
                # Find the corresponding table object
                for table in doc.tables:
                    if table._element == element:
                        markdown_table = self.convert_table_to_markdown(table)
                        markdown_parts.append(markdown_table)
                        break
        
        return "".join(markdown_parts)

def convert_docx_to_markdown(input_path, output_path):
    """
    Convert a DOCX file to Markdown format.
    
    Args:
        input_path (str): Path to the input DOCX file
        output_path (str): Path for the output Markdown file
    """
    try:
        # Load the document
        print(f"Loading document: {input_path}")
        doc = Document(input_path)
        
        # Create converter instance
        converter = DocxToMarkdownConverter()
        
        # Convert document to markdown
        print("Converting to Markdown...")
        markdown_content = converter.convert_document(doc)
        
        # Clean up extra whitespace
        markdown_content = re.sub(r'\n{3,}', '\n\n', markdown_content)
        markdown_content = markdown_content.strip()
        
        # Save the markdown file
        print(f"Saving Markdown file: {output_path}")
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        print("Conversion completed successfully!")
        return True
        
    except FileNotFoundError:
        print(f"Error: Input file '{input_path}' not found.")
        return False
    except PermissionError:
        print(f"Error: Permission denied. Check file permissions.")
        return False
    except Exception as e:
        print(f"Error converting document: {str(e)}")
        return False

def main():
    """
    Main function with hardcoded file paths.
    """
    # ===== CONFIGURE YOUR FILE PATHS HERE =====
    input_file = "D:\\Downloads\\code\\code\\Extra-Projects\\rag-pipeline-eqanun\\document_formatting\\mecelleler-docx\\output-cinayet.docx"     # Change this to your input DOCX file path
    output_file = "output.md"     # Change this to your desired output Markdown file path
    # ===========================================
    
    # Validate input file exists
    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' does not exist.")
        return
    
    # Validate input file extension
    if not input_file.lower().endswith('.docx'):
        print("Warning: Input file doesn't have .docx extension. Proceeding anyway...")
    
    # Ensure output file has .md extension
    if not output_file.lower().endswith('.md'):
        base_name = os.path.splitext(output_file)[0]
        output_file = base_name + '.md'
        print(f"Output file extension changed to: {output_file}")
    
    # Convert the document
    success = convert_docx_to_markdown(input_file, output_file)
    
    if success:
        print(f"\nSuccessfully converted '{input_file}' -> '{output_file}'")
        print("The DOCX file has been converted to Markdown format.")
        
        # Show file size info
        try:
            input_size = os.path.getsize(input_file)
            output_size = os.path.getsize(output_file)
            print(f"Original size: {input_size:,} bytes")
            print(f"Markdown size: {output_size:,} bytes")
        except:
            pass
    else:
        print("Failed to convert the document.")

if __name__ == "__main__":
    main()